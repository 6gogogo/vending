#!/usr/bin/env python3
"""从四份 Markdown 正文生成紧凑版 Word 用户手册。"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass, field
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
MANUAL_DIR = REPO_ROOT / "docs" / "用户手册"
ROLE_FILES = {
    "实例管理员": ("实例管理员使用手册.md", "公益智助柜实例管理员使用手册.docx"),
    "商户": ("商户使用手册.md", "公益智助柜商户使用手册.docx"),
    "补货员": ("补货员使用手册.md", "公益智助柜补货员使用手册.docx"),
    "App 用户": ("App用户使用手册.md", "公益智助柜App用户使用手册.docx"),
}

GREEN = "087A63"
DEEP_GREEN = "075C4D"
PALE_GREEN = "EDF7F1"
PALE_BLUE = "EEF5F8"
PALE_ORANGE = "FFF7EA"
INK = RGBColor(18, 39, 55)
MUTED = RGBColor(88, 108, 124)
CONTENT_WIDTH_DXA = 10028
TABLE_INDENT_DXA = 120


@dataclass
class ManualSection:
    title: str
    purpose: str = ""
    steps: list[str] = field(default_factory=list)
    images: list[tuple[str, Path]] = field(default_factory=list)
    completion: str = ""
    issues: list[tuple[str, str]] = field(default_factory=list)


@dataclass
class Manual:
    title: str
    intro: str
    sections: list[ManualSection]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = GREEN, width: str = "10") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), width)
        element.set(qn("w:color"), color)


def set_paragraph_box(paragraph, fill: str, color: str = GREEN, width: str = "10", edges: tuple[str, ...] = ("top", "left", "bottom", "right")) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge in edges:
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), width)
        element.set(qn("w:space"), "7")
        element.set(qn("w:color"), color)
        borders.append(element)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    width = table_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        table_pr.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")

    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        grid.append(grid_column)

    for row in table.rows:
        for cell, column_width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_width = tc_pr.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                tc_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(column_width))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def create_step_numbering(document: Document) -> int:
    numbering = document.part.numbering_part._element
    existing_abstract = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_abstract, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    for tag, value in (("start", "1"), ("numFmt", "decimal"), ("lvlText", "%1."), ("lvlJc", "left")):
        element = OxmlElement(f"w:{tag}")
        element.set(qn("w:val"), value)
        level.append(element)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Microsoft YaHei")
    fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    fonts.set(qn("w:eastAsia"), "微软雅黑")
    r_pr.append(fonts)
    bold = OxmlElement("w:b")
    r_pr.append(bold)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN)
    r_pr.append(color)
    level.append(r_pr)
    abstract.append(level)
    numbering.insert(0, abstract)
    return abstract_id


def create_numbering_instance(document: Document, abstract_id: int) -> int:
    numbering = document.part.numbering_part._element
    existing_nums = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(existing_nums, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(num_ref)
    p_pr.append(num_pr)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """为段落添加内部跳转书签。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph_properties = paragraph._p.find(qn("w:pPr"))
    insert_at = 1 if paragraph_properties is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text: str, anchor: str) -> None:
    """添加不依赖外部关系的 Word 文档内链接。"""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = paragraph.add_run(text)
    set_font(run, 10.5, bold=True, color=RGBColor(0, 108, 86))
    run.font.underline = False
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)


def set_font(run, size: float, bold: bool = False, color: RGBColor = INK) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_font(run, 8.5, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    tail = paragraph.add_run(" 页")
    set_font(tail, 8.5, color=MUTED)


def parse_manual(path: Path) -> Manual:
    lines = path.read_text(encoding="utf-8").splitlines()
    title = ""
    intro_lines: list[str] = []
    sections: list[ManualSection] = []
    current: ManualSection | None = None

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if line.startswith("## "):
            current = ManualSection(title=line[3:].strip())
            sections.append(current)
            continue
        if current is None:
            intro_lines.append(line)
            continue
        if line.startswith("**用途：**"):
            current.purpose = line.removeprefix("**用途：**").strip()
            continue
        if line.startswith("**完成标志：**"):
            current.completion = line.removeprefix("**完成标志：**").strip()
            continue
        step_match = re.match(r"^\d+\.\s+(.+)$", line)
        if step_match:
            current.steps.append(step_match.group(1).strip())
            continue
        image_match = re.match(r"^!\[(.+?)\]\((.+?)\)$", line)
        if image_match:
            caption, relative = image_match.groups()
            current.images.append((caption, (path.parent / relative).resolve()))
            continue
        issue_match = re.match(r"^-\s+\*\*(.+?)：\*\*\s*(.+)$", line)
        if issue_match:
            current.issues.append((issue_match.group(1), issue_match.group(2)))

    if not title or not sections:
        raise ValueError(f"手册结构不完整：{path}")
    return Manual(title=title, intro=" ".join(intro_lines), sections=sections)


def configure_document(document: Document, title: str) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in (("Title", 20, INK), ("Heading 1", 15, INK), ("Heading 2", 13, INK)):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    if "Figure Caption" not in document.styles:
        caption_style = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption_style = document.styles["Figure Caption"]
    caption_style.font.name = "Microsoft YaHei"
    caption_style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    caption_style.font.size = Pt(8.5)
    caption_style.font.color.rgb = MUTED
    caption_style.paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"公益智助柜　{title.split('｜')[-1]}")
    set_font(run, 8.5, bold=True, color=RGBColor(50, 111, 94))
    add_page_number(section.footer.paragraphs[0])


def add_title_block(document: Document, manual: Manual) -> None:
    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_before = Pt(7)
    title_paragraph.paragraph_format.space_after = Pt(3)
    title_paragraph.paragraph_format.keep_with_next = True
    set_paragraph_box(title_paragraph, PALE_GREEN, width="14", edges=("top", "left", "right"))
    title_run = title_paragraph.add_run(manual.title)
    set_font(title_run, 19, bold=True)
    intro_paragraph = document.add_paragraph()
    intro_paragraph.paragraph_format.space_before = Pt(0)
    intro_paragraph.paragraph_format.space_after = Pt(7)
    set_paragraph_box(intro_paragraph, PALE_GREEN, width="14", edges=("left", "bottom", "right"))
    intro_run = intro_paragraph.add_run(manual.intro)
    set_font(intro_run, 10.5, color=MUTED)

    meta = document.add_paragraph()
    meta.paragraph_format.space_before = Pt(4)
    meta.paragraph_format.space_after = Pt(10)
    meta_run = meta.add_run("版本日期：2026 年 8 月 9 日　｜　按任务查找，图中编号对应操作顺序")
    set_font(meta_run, 8.5, color=MUTED)


def add_toc(document: Document, manual: Manual) -> list[str]:
    """在标题区后生成静态、可点击的任务目录。"""
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.space_before = Pt(2)
    heading.paragraph_format.space_after = Pt(3)
    set_keep_with_next(heading)
    run = heading.add_run("目录")
    set_font(run, 15, bold=True)
    add_bookmark(heading, "manual_toc", 1)

    hint = document.add_paragraph()
    hint.paragraph_format.space_before = Pt(0)
    hint.paragraph_format.space_after = Pt(5)
    hint_run = hint.add_run("点击目录条目可跳转到对应任务。")
    set_font(hint_run, 9, color=MUTED)

    bookmark_names: list[str] = []
    for index, section in enumerate(manual.sections, start=1):
        bookmark_name = f"manual_section_{index:02d}"
        bookmark_names.append(bookmark_name)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.35)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        add_internal_link(paragraph, section.title, bookmark_name)

    document.add_page_break()
    return bookmark_names


def add_info_box(document: Document, label: str, text: str, fill: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    set_paragraph_box(paragraph, fill, width="7")
    label_run = paragraph.add_run(f"{label}　")
    set_font(label_run, 10.2, bold=True, color=RGBColor(0, 108, 86))
    text_run = paragraph.add_run(text)
    set_font(text_run, 10.2)


def add_screenshot(document: Document, caption: str, image_path: Path) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"正文引用的截图不存在：{image_path}")
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    max_width_cm = 16.7
    max_height_cm = 12.6
    aspect = width_px / height_px
    width_cm = min(max_width_cm, max_height_cm * aspect)
    height_cm = width_cm / aspect
    if width_px <= 500:
        height_cm = min(12.9, max_height_cm)
        width_cm = height_cm * aspect

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    shape = paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    caption_paragraph = document.add_paragraph(caption, style="Figure Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_screenshot_pair(document: Document, images: list[tuple[str, Path]]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2)
    for index, (_, image_path) in enumerate(images):
        if not image_path.exists():
            raise FileNotFoundError(f"正文引用的截图不存在：{image_path}")
        with Image.open(image_path) as image:
            width_px, height_px = image.size
        aspect = width_px / height_px
        width_cm = min(7.9, 14.0 * aspect)
        height_cm = width_cm / aspect
        if index:
            spacer = paragraph.add_run("　")
            set_font(spacer, 8.5, color=MUTED)
        caption = images[index][0]
        shape = paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm), height=Cm(height_cm))
        shape._inline.docPr.set("descr", caption)
        shape._inline.docPr.set("title", caption)

    captions = "　｜　".join(caption for caption, _ in images)
    caption_paragraph = document.add_paragraph(captions, style="Figure Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_task_section(
    document: Document,
    section: ManualSection,
    abstract_num_id: int,
    bookmark_name: str,
    bookmark_id: int,
    page_break_before: bool,
) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.page_break_before = page_break_before
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(5)
    set_keep_with_next(heading)
    run = heading.add_run(section.title)
    set_font(run, 15, bold=True)
    add_bookmark(heading, bookmark_name, bookmark_id)

    add_info_box(document, "用途", section.purpose, PALE_BLUE)
    steps_heading = document.add_paragraph()
    steps_heading.paragraph_format.space_before = Pt(6)
    steps_heading.paragraph_format.space_after = Pt(3)
    set_keep_with_next(steps_heading)
    run = steps_heading.add_run("操作步骤")
    set_font(run, 10.5, bold=True, color=RGBColor(0, 108, 86))

    num_id = create_numbering_instance(document, abstract_num_id)
    for step in section.steps:
        paragraph = document.add_paragraph(style="Normal")
        apply_numbering(paragraph, num_id)
        text_run = paragraph.add_run(step)
        set_font(text_run, 10.5)

    if len(section.images) == 2:
        add_screenshot_pair(document, section.images)
    else:
        for caption, image_path in section.images:
            add_screenshot(document, caption, image_path)
    add_info_box(document, "完成标志", section.completion, PALE_GREEN)


def add_issue_page(document: Document, section: ManualSection, bookmark_name: str, bookmark_id: int) -> None:
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.page_break_before = True
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run("遇到问题")
    set_font(run, 16, bold=True)
    add_bookmark(heading, bookmark_name, bookmark_id)

    intro = document.add_paragraph()
    intro_run = intro.add_run("只保留会阻塞当前操作的问题。先按对应处理方式核对，再继续原任务。")
    set_font(intro_run, 10.5, color=MUTED)

    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [2950, 7078])
    headers = table.rows[0].cells
    headers[0].text = "现象"
    headers[1].text = "处理方式"
    set_repeat_table_header(table.rows[0])
    for cell in headers:
        set_cell_shading(cell, DEEP_GREEN)
        set_cell_border(cell, DEEP_GREEN, "8")
        for run in cell.paragraphs[0].runs:
            set_font(run, 10.2, bold=True, color=RGBColor(255, 255, 255))

    for index, (problem, resolution) in enumerate(section.issues):
        cells = table.add_row().cells
        set_table_geometry(table, [2950, 7078])
        cells[0].text = problem
        cells[1].text = resolution
        fill = PALE_ORANGE if index % 2 == 0 else "FFFFFF"
        for cell in cells:
            set_cell_shading(cell, fill)
            set_cell_border(cell, "D7E6DF", "6")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_font(run, 9.7, bold=(cell is cells[0]))


def build_manual(role: str, source: Path, output: Path) -> None:
    manual = parse_manual(source)
    document = Document()
    configure_document(document, manual.title)
    abstract_num_id = create_step_numbering(document)
    add_title_block(document, manual)
    bookmark_names = add_toc(document, manual)
    task_index = 0
    for section_index, section in enumerate(manual.sections):
        bookmark_name = bookmark_names[section_index]
        bookmark_id = section_index + 2
        if section.title == "遇到问题":
            add_issue_page(document, section, bookmark_name, bookmark_id)
        else:
            add_task_section(
                document,
                section,
                abstract_num_id,
                bookmark_name,
                bookmark_id,
                page_break_before=task_index > 0,
            )
            task_index += 1

    document.core_properties.title = manual.title
    document.core_properties.subject = f"公益智助柜{role}操作说明"
    document.core_properties.author = "公益智助柜"
    document.core_properties.keywords = "公益智助柜, 用户手册, 操作说明"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"已生成：{output.relative_to(REPO_ROOT)}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--role", choices=list(ROLE_FILES), help="仅生成指定身份手册")
    args = parser.parse_args()

    roles = [args.role] if args.role else list(ROLE_FILES)
    for role in roles:
        source_name, output_name = ROLE_FILES[role]
        build_manual(role, MANUAL_DIR / source_name, MANUAL_DIR / output_name)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
